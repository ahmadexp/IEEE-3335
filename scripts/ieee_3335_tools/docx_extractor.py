"""DOCX extraction and conversion utilities."""

from pathlib import Path
from typing import Optional

# Import shared console from package
from . import console


def extract_docx_to_markdown(
    docx_path: Path, 
    output_path: Optional[Path] = None,
    extract_images: bool = True,
    quiet: bool = False
) -> Path:
    """
    Extract DOCX to markdown using python-docx.
    
    Args:
        docx_path: Path to input DOCX file
        output_path: Optional output path (default: {input_name}.md)
        extract_images: Whether to extract and save images (currently not implemented for DOCX)
        quiet: Whether to suppress individual file output messages
        
    Returns:
        Path to the created markdown file
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is not installed. Please install it with: uv add python-docx"
        )
    
    if output_path is None:
        output_path = docx_path.with_suffix('.md')
    
    if not quiet:
        console.print(f"[bold]Converting DOCX to Markdown:[/bold] {docx_path}")
    
    doc = Document(str(docx_path))
    markdown_lines = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            markdown_lines.append("")
            continue
        
        # Check for heading styles
        if para.style.name.startswith('Heading'):
            try:
                level = int(para.style.name.split()[-1])
                markdown_lines.append(f"{'#' * level} {text}")
            except (ValueError, IndexError):
                markdown_lines.append(f"## {text}")
        else:
            # Check for bold/italic formatting
            formatted_text = ""
            for run in para.runs:
                run_text = run.text
                if run.bold and run.italic:
                    formatted_text += f"***{run_text}***"
                elif run.bold:
                    formatted_text += f"**{run_text}**"
                elif run.italic:
                    formatted_text += f"*{run_text}*"
                else:
                    formatted_text += run_text
            markdown_lines.append(formatted_text)
    
    # Handle tables
    for table in doc.tables:
        markdown_lines.append("")
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            markdown_lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                markdown_lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
        markdown_lines.append("")
    
    # Write output
    markdown_content = "\n".join(markdown_lines)
    output_path.write_text(markdown_content, encoding='utf-8')
    
    if not quiet:
        console.print(f"[green]✅ Successfully converted to markdown:[/green] {output_path}")
    return output_path
